from __future__ import annotations

import io
import re
from pathlib import Path


def _extract_pdf(data: bytes) -> tuple[str, dict]:
    import pymupdf  # type: ignore

    doc = pymupdf.open(stream=data, filetype="pdf")
    text = "\n\n".join(page.get_text() for page in doc)
    meta = doc.metadata or {}
    citation: dict = {}
    if meta.get("author"):
        citation["author"] = meta["author"]
    if meta.get("title"):
        citation["title"] = meta["title"]
    return text, citation


def _extract_docx(data: bytes) -> tuple[str, dict]:
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(data))
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    props = doc.core_properties
    citation: dict = {}
    if props.author:
        citation["author"] = props.author
    if props.title:
        citation["title"] = props.title
    return text, citation


def _extract_html(html: str) -> tuple[str, dict]:
    import trafilatura  # type: ignore
    from trafilatura.metadata import extract_metadata  # type: ignore

    text = trafilatura.extract(html) or ""
    citation: dict = {}
    try:
        meta = extract_metadata(html)
        if meta:
            if getattr(meta, "title", None):
                citation["title"] = meta.title
            if getattr(meta, "author", None):
                citation["author"] = meta.author
            if getattr(meta, "sitename", None):
                citation["site_name"] = meta.sitename
            if getattr(meta, "date", None):
                citation["pub_date"] = meta.date
            if getattr(meta, "url", None):
                citation["url"] = meta.url
    except Exception:
        pass
    return text, citation


def extract_file(data: bytes, filename: str) -> tuple[str, dict]:
    """Return (text, citation_metadata) from raw file bytes."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    # TXT / plain text
    text = data.decode("utf-8", errors="replace")
    return text, {}


def extract_file_pages(content: bytes, filename: str):
    """Yield (page_number, text) tuples; generator so large PDFs stream page by page."""
    name = filename.lower()
    if name.endswith(".pdf"):
        yield from _extract_pdf_pages(content)
    elif name.endswith(".docx"):
        yield from _extract_docx_pages(content)
    else:
        yield (1, content.decode("utf-8", errors="replace"))


def _extract_pdf_pages(data: bytes):
    import pymupdf  # type: ignore

    doc = pymupdf.open(stream=data, filetype="pdf")
    for i, page in enumerate(doc):
        yield (i + 1, page.get_text())


def _extract_docx_pages(data: bytes):
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(data))
    sections: list[list[str]] = []
    current: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style.name.lower().startswith("heading")
        if is_heading:
            if current:
                sections.append(current)
            current = [text]
        else:
            current.append(text)

    if current:
        sections.append(current)

    if not sections:
        yield (1, "")
        return

    for i, lines in enumerate(sections):
        yield (i + 1, "\n".join(lines))


def extract_url(html: str) -> tuple[str, dict]:
    """Return (text, citation_metadata) from fetched HTML."""
    return _extract_html(html)
